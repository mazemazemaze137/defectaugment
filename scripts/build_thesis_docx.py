from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / "毕业论文_初稿.md"
OUTPUT_PATH = ROOT / "毕业论文_初稿_格式化.docx"

TITLE = "基于生成对抗网络的工业表面缺陷数据增强系统设计与实现"


def set_run_font(run, size: float | None = None, bold: bool | None = None, name: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_style_font(style, size: float, east_asia: str = "宋体", bold: bool = False) -> None:
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def set_paragraph_body(paragraph, first_line: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(24)


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text.strip()


def add_body_paragraph(doc: Document, text: str, first_line: bool = True, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    else:
        set_paragraph_body(p, first_line=first_line)
    run = p.add_run(clean_inline(text))
    set_run_font(run, 12)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
    run = p.add_run(clean_inline(text))
    set_run_font(run, 16 if level == 1 else 14, bold=True, name="黑体")


def add_keywords(doc: Document, line: str) -> None:
    text = clean_inline(line)
    if text.startswith("Keywords"):
        label, rest = text.split(":", 1)
        label += ":"
    elif text.startswith("关键词"):
        label, rest = text.split("：", 1)
        label += "："
    elif text.startswith("关键字"):
        label, rest = text.split("：", 1)
        label += "："
    elif "：" in text:
        label, rest = text.split("：", 1)
        label += "："
    elif ":" in text:
        label, rest = text.split(":", 1)
        label += ":"
    else:
        label, rest = "关键词：", text
    p = doc.add_paragraph()
    set_paragraph_body(p, first_line=False)
    if label.startswith("Keywords"):
        label = "Keywords: "
    r1 = p.add_run(label)
    set_run_font(r1, 12, bold=True, name="黑体")
    r2 = p.add_run(rest.strip())
    set_run_font(r2, 12)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean_inline(value))
            set_run_font(run, 10.5, bold=(i == 0), name="黑体" if i == 0 else "宋体")
    doc.add_paragraph()


def add_image(doc: Document, caption: str, rel_path: str, fig_no: int) -> int:
    image_path = ROOT / rel_path
    if not image_path.exists():
        add_body_paragraph(doc, f"（图像文件未找到：{rel_path}）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        return fig_no
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(f"图 {fig_no} {caption}")
    set_run_font(r, 10.5)
    return fig_no + 1


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    rows = []
    for line in raw:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells):
            continue
        rows.append(cells)
    return rows, i


def setup_document() -> Document:
    doc = Document()
    configure_section(doc.sections[0])

    styles = doc.styles
    set_style_font(styles["Normal"], 12, "宋体")
    set_style_font(styles["Heading 1"], 16, "黑体", True)
    set_style_font(styles["Heading 2"], 14, "黑体", True)
    set_style_font(styles["Heading 3"], 12, "黑体", True)
    set_style_font(styles["Title"], 22, "黑体", True)
    return doc


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)


def add_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, 10.5)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), "1")
    section._sectPr.append(pg_num)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本科毕业设计（论文）")
    set_run_font(r, 22, bold=True, name="黑体")

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run_font(r, 18, bold=True, name="黑体")

    for _ in range(5):
        doc.add_paragraph()
    meta = [
        ("学院", "（待填写）"),
        ("专业", "（待填写）"),
        ("班级", "（待填写）"),
        ("学生姓名", "（待填写）"),
        ("学号", "（待填写）"),
        ("指导教师", "（待填写）"),
        ("完成日期", "2026 年 4 月"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：{value}")
        set_run_font(r, 14)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    add_page_number(section)


def add_toc(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("目 录")
    set_run_font(r, 16, bold=True, name="黑体")
    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    doc.add_page_break()


def build() -> None:
    lines = MARKDOWN_PATH.read_text(encoding="utf-8").splitlines()
    doc = setup_document()
    add_cover(doc)

    fig_no = 1
    i = 0
    saw_first_chapter = False
    pending_toc = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            heading = clean_inline(line[3:])
            if heading in {"Abstract", "参考文献", "致谢"}:
                doc.add_page_break()
            if heading == "第一章 绪论" and not pending_toc:
                add_toc(doc)
                pending_toc = True
            if heading.startswith("第") or heading in {"摘要", "Abstract", "参考文献", "致谢"}:
                if heading.startswith("第") and saw_first_chapter:
                    doc.add_page_break()
                if heading.startswith("第一章"):
                    saw_first_chapter = True
                add_heading(doc, heading, 1)
            else:
                add_heading(doc, heading, 1)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:], 2)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            fig_no = add_image(doc, image_match.group(1), image_match.group(2), fig_no)
            i += 1
            continue
        if line.startswith("**关键词") or line.startswith("**Keywords"):
            add_keywords(doc, line)
            i += 1
            continue
        if re.match(r"\[\d+\]", line):
            add_body_paragraph(doc, line, first_line=False)
            i += 1
            continue
        add_body_paragraph(doc, line)
        i += 1

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build()
    print(OUTPUT_PATH)
